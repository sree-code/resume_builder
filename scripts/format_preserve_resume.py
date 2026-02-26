#!/usr/bin/env python3
import argparse
from copy import deepcopy
import json
import sys
from pathlib import Path


def eprint(*args):
    print(*args, file=sys.stderr)


def _load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _save_json(path, data):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _clean_line_text(text):
    return (text or "").replace("\r", " ").replace("\n", " ").strip()


def flatten_docx_paragraphs(doc):
    from docx.document import Document as _Document
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    def iter_block_items(parent):
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            return
        for child in parent_elm.iterchildren():
            tag = child.tag.rsplit("}", 1)[-1]
            if tag == "p":
                yield Paragraph(child, parent)
            elif tag == "tbl":
                yield Table(child, parent)

    paragraphs = []
    seen_cells = set()

    def walk(parent):
        for item in iter_block_items(parent):
            if item.__class__.__name__ == "Paragraph":
                paragraphs.append(item)
            else:
                for row in item.rows:
                    for cell in row.cells:
                        cell_id = id(cell._tc)
                        if cell_id in seen_cells:
                            continue
                        seen_cells.add(cell_id)
                        walk(cell)

    walk(doc)
    return paragraphs


def extract_docx(input_path):
    from docx import Document

    doc = Document(str(input_path))
    paragraphs = flatten_docx_paragraphs(doc)
    lines = []
    line_mappings = []
    for idx, para in enumerate(paragraphs, start=1):
        text = _clean_line_text(para.text)
        lines.append(text)
        line_mappings.append(
            {
                "lineNumber": idx,
                "target": "docx_paragraph",
                "paragraphIndex": idx - 1,
                "isEmpty": text == "",
            }
        )

    return {
        "kind": "docx",
        "text": "\n".join(lines),
        "lineMappings": line_mappings,
        "meta": {
            "paragraphCount": len(paragraphs),
            "lineCount": len(lines),
        },
    }


def _set_paragraph_text_preserve_para_style(paragraph, new_text):
    runs = list(paragraph.runs)
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def _normalize_docx_insert_text_for_anchor(anchor_paragraph, new_text):
    text = _clean_line_text(new_text)
    if not text:
        return text
    anchor_text = _clean_line_text(getattr(anchor_paragraph, "text", ""))
    # If the anchor paragraph is a Word list bullet (visible bullet not in text),
    # avoid storing a literal "- " prefix because Word will render its own bullet.
    if anchor_text and not anchor_text.lstrip().startswith(("-", "*", "•")):
        text = text.lstrip()
        text = text[2:] if text.startswith("- ") else text
        text = text[2:] if text.startswith("* ") else text
        text = text[2:] if text.startswith("• ") else text
    return text


def _insert_docx_paragraph_after(anchor_paragraph, text):
    from docx.text.paragraph import Paragraph

    new_p = deepcopy(anchor_paragraph._p)
    anchor_paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, anchor_paragraph._parent)
    _set_paragraph_text_preserve_para_style(new_para, text)
    return new_para


def apply_docx(input_path, output_path, mapping_payload, line_edits, insertions=None):
    from docx import Document

    doc = Document(str(input_path))
    paragraphs = flatten_docx_paragraphs(doc)
    mapping_by_line = {m["lineNumber"]: m for m in mapping_payload.get("lineMappings", []) if m.get("target") == "docx_paragraph"}

    applied = []
    for edit in line_edits or []:
        line_number = edit.get("lineNumber")
        new_text = edit.get("newText")
        if not isinstance(line_number, int) or not isinstance(new_text, str):
            continue
        mapping = mapping_by_line.get(line_number)
        if not mapping:
            continue
        para_idx = mapping.get("paragraphIndex")
        if not isinstance(para_idx, int) or para_idx < 0 or para_idx >= len(paragraphs):
            continue
        para = paragraphs[para_idx]
        _set_paragraph_text_preserve_para_style(para, new_text)
        applied.append({"lineNumber": line_number, "paragraphIndex": para_idx})

    inserted = []
    normalized_insertions = []
    for order_idx, ins in enumerate(insertions or []):
        after_line = ins.get("afterLineNumber")
        new_text = ins.get("newText")
        if not isinstance(after_line, int) or not isinstance(new_text, str):
            continue
        mapping = mapping_by_line.get(after_line)
        if not mapping:
            continue
        para_idx = mapping.get("paragraphIndex")
        if not isinstance(para_idx, int) or para_idx < 0 or para_idx >= len(paragraphs):
            continue
        normalized_insertions.append(
            {
                "afterLineNumber": after_line,
                "paragraphIndex": para_idx,
                "newText": new_text,
                "order": order_idx,
            }
        )

    # Insert in reverse order per anchor to preserve requested sequence after the same line.
    for ins in sorted(normalized_insertions, key=lambda x: (x["paragraphIndex"], x["order"]), reverse=True):
        anchor_para = paragraphs[ins["paragraphIndex"]]
        insert_text = _normalize_docx_insert_text_for_anchor(anchor_para, ins["newText"])
        if not insert_text:
            continue
        _insert_docx_paragraph_after(anchor_para, insert_text)
        inserted.append(
            {
                "afterLineNumber": ins["afterLineNumber"],
                "paragraphIndex": ins["paragraphIndex"],
            }
        )

    doc.save(str(output_path))
    return {"applied": applied, "inserted": inserted, "outputPath": str(output_path)}


def _color_int_to_rgb_tuple(color_int):
    if not isinstance(color_int, int):
        return (0, 0, 0)
    r = (color_int >> 16) & 255
    g = (color_int >> 8) & 255
    b = color_int & 255
    return (r / 255.0, g / 255.0, b / 255.0)


def extract_pdf(input_path):
    import fitz

    doc = fitz.open(str(input_path))
    lines = []
    line_mappings = []
    line_no = 1

    for page_index, page in enumerate(doc):
        data = page.get_text("dict")
        for block in data.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                text = "".join(span.get("text", "") for span in spans)
                cleaned = _clean_line_text(text)
                if cleaned == "":
                    continue
                first = spans[0] if spans else {}
                bbox = line.get("bbox") or block.get("bbox")
                lines.append(cleaned)
                line_mappings.append(
                    {
                        "lineNumber": line_no,
                        "target": "pdf_line",
                        "pageIndex": page_index,
                        "bbox": list(bbox) if bbox else None,
                        "fontSize": float(first.get("size", 10)),
                        "fontName": str(first.get("font", "")),
                        "color": int(first.get("color", 0)),
                    }
                )
                line_no += 1

    return {
        "kind": "pdf",
        "text": "\n".join(lines),
        "lineMappings": line_mappings,
        "meta": {"pageCount": doc.page_count, "lineCount": len(lines)},
    }


def _insert_pdf_line(page, rect, text, font_size, color_int):
    import fitz

    color = _color_int_to_rgb_tuple(color_int)
    size = max(6, min(18, float(font_size or 10)))
    for _ in range(8):
        spare = page.insert_textbox(
            rect,
            text,
            fontname="helv",
            fontsize=size,
            color=color,
            align=fitz.TEXT_ALIGN_LEFT,
        )
        if spare >= 0:
            return {"fontSizeUsed": size, "fit": True}
        size = max(5, size - 0.7)
        # Clear overlay attempts by repainting white before trying smaller size.
        page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
    page.insert_textbox(
        rect,
        text,
        fontname="helv",
        fontsize=size,
        color=color,
        align=fitz.TEXT_ALIGN_LEFT,
    )
    return {"fontSizeUsed": size, "fit": False}


def apply_pdf(input_path, output_path, mapping_payload, edits_payload):
    import fitz

    doc = fitz.open(str(input_path))
    mapping_by_line = {m["lineNumber"]: m for m in mapping_payload.get("lineMappings", []) if m.get("target") == "pdf_line"}

    edits_by_page = {}
    for edit in edits_payload:
        line_number = edit.get("lineNumber")
        new_text = edit.get("newText")
        if not isinstance(line_number, int) or not isinstance(new_text, str):
            continue
        mapping = mapping_by_line.get(line_number)
        if not mapping or not mapping.get("bbox"):
            continue
        page_idx = mapping.get("pageIndex")
        if not isinstance(page_idx, int) or page_idx < 0 or page_idx >= doc.page_count:
            continue
        edits_by_page.setdefault(page_idx, []).append((mapping, new_text))

    applied = []
    for page_idx, page_edits in edits_by_page.items():
        page = doc[page_idx]
        # Redact original line areas first, then overlay replacement text.
        for mapping, _new_text in page_edits:
            rect = fitz.Rect(mapping["bbox"])
            page.add_redact_annot(rect, fill=(1, 1, 1))
        page.apply_redactions()

        for mapping, new_text in page_edits:
            rect = fitz.Rect(mapping["bbox"])
            result = _insert_pdf_line(page, rect, new_text, mapping.get("fontSize", 10), mapping.get("color", 0))
            applied.append(
                {
                    "lineNumber": mapping["lineNumber"],
                    "pageIndex": page_idx,
                    "fontSizeUsed": result["fontSizeUsed"],
                    "fit": result["fit"],
                }
            )

    doc.save(str(output_path), garbage=4, deflate=True)
    return {"applied": applied, "outputPath": str(output_path)}


def cmd_extract(args):
    input_path = Path(args.input)
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        payload = extract_docx(input_path)
    elif suffix == ".pdf":
        payload = extract_pdf(input_path)
    else:
        raise ValueError("Supported input types: .docx, .pdf")
    _save_json(args.output, payload)
    print(json.dumps({"ok": True, "kind": payload["kind"], "meta": payload.get("meta", {})}))


def cmd_apply(args):
    input_path = Path(args.input)
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    mapping_payload = _load_json(args.mapping)
    edits_payload = _load_json(args.edits)
    if isinstance(edits_payload, dict):
        line_edits = edits_payload.get("lineEdits", [])
        insertions = edits_payload.get("insertions", [])
    else:
        line_edits = edits_payload
        insertions = []
    suffix = input_path.suffix.lower()

    if suffix == ".docx":
        result = apply_docx(input_path, output_path, mapping_payload, line_edits, insertions)
    elif suffix == ".pdf":
        result = apply_pdf(input_path, output_path, mapping_payload, line_edits)
    else:
        raise ValueError("Supported input types: .docx, .pdf")
    print(
        json.dumps(
            {
                "ok": True,
                "appliedCount": len(result.get("applied", [])),
                "insertedCount": len(result.get("inserted", [])) if isinstance(result, dict) else 0,
                "outputPath": str(output_path),
            }
        )
    )


def main():
    parser = argparse.ArgumentParser(description="Format-preserving resume extract/apply for DOCX/PDF")
    sub = parser.add_subparsers(dest="command", required=True)

    p_extract = sub.add_parser("extract")
    p_extract.add_argument("--input", required=True)
    p_extract.add_argument("--output", required=True)
    p_extract.set_defaults(func=cmd_extract)

    p_apply = sub.add_parser("apply")
    p_apply.add_argument("--input", required=True)
    p_apply.add_argument("--mapping", required=True)
    p_apply.add_argument("--edits", required=True)
    p_apply.add_argument("--output", required=True)
    p_apply.set_defaults(func=cmd_apply)

    args = parser.parse_args()
    try:
        args.func(args)
    except Exception as exc:
        eprint(f"ERROR: {exc}")
        sys.exit(1)


if __name__ == "__main__":
    main()
